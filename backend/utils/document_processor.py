import io
from typing import Dict, Optional
import pdfplumber
import PyPDF2
from docx import Document
import logging
import traceback
import re

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(message)s',
    force=True
)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF file bytes"""
        text_parts = []
        errors = []
        
        # Try pdfplumber first (usually better quality)
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text(
                        x_tolerance=2,  # Adjust for better word spacing
                        y_tolerance=2,  # Adjust for better line spacing
                    )
                    if text:
                        # Clean the text
                        text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces with single space
                        text = text.strip()
                        if text:
                            text_parts.append(text)
                
                if text_parts:
                    logger.info(f"Successfully extracted text from PDF using pdfplumber: {len(text_parts)} pages")
        except Exception as e:
            error_msg = f"Error extracting text with pdfplumber: {str(e)}"
            logger.warning(error_msg)
            errors.append(error_msg)
        
        # If pdfplumber failed, try PyPDF2 as fallback
        if not text_parts:
            try:
                pdf_file = io.BytesIO(file_bytes)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text = text.strip()
                        if text:
                            text_parts.append(text)
                
                if text_parts:
                    logger.info(f"Successfully extracted text from PDF using PyPDF2: {len(text_parts)} pages")
            except Exception as e:
                error_msg = f"Error extracting text with PyPDF2: {str(e)}"
                logger.warning(error_msg)
                errors.append(error_msg)
        
        # If still no text, raise error
        if not text_parts:
            error_details = "; ".join(errors)
            raise ValueError(f"Failed to extract text from PDF: {error_details}")
        
        # Join all parts with newlines
        full_text = '\n'.join(text_parts)
        
        # Clean up the text
        full_text = DocumentProcessor._clean_text(full_text)
        
        logger.info(f"Extracted {len(full_text)} characters from PDF")
        return full_text
    
    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX file bytes"""
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)
            
            # Extract text from paragraphs
            text_parts = []
            
            # Get text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            # Join with newlines and clean
            full_text = '\n'.join(text_parts)
            if not full_text.strip():
                raise ValueError("No text content found in the DOCX file")
                
            full_text = DocumentProcessor._clean_text(full_text)
            
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            logger.error(traceback.format_exc())
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
    
    @staticmethod
    def process_resume(file_bytes: bytes, filename: str) -> str:
        """Process resume file and extract text"""
        try:
            if not file_bytes:
                raise ValueError("Empty file received")
            
            # Extract text based on file type
            filename_lower = filename.lower()
            if filename_lower.endswith('.pdf'):
                text = DocumentProcessor.extract_text_from_pdf(file_bytes)
            elif filename_lower.endswith(('.doc', '.docx')):
                text = DocumentProcessor.extract_text_from_docx(file_bytes)
            else:
                raise ValueError("Unsupported file type. Only PDF and Word documents are supported.")
            
            if not text:
                raise ValueError("No text could be extracted from the file")
            
            return text
            
        except Exception as e:
            raise ValueError(f"Error processing resume: {str(e)}")
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean extracted text"""
        if not text:
            return ""
        
        # Replace Unicode quotes and apostrophes with ASCII versions
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        # Replace Unicode spaces and separators with regular spaces
        text = re.sub(r'[\u00A0\u1680\u2000-\u200A\u2028\u2029\u202F\u205F\u3000]', ' ', text)
        
        # Replace multiple spaces with single space
        text = re.sub(r'\s+', ' ', text)
        
        # Add space after periods if missing
        text = re.sub(r'\.(?=[A-Z])', '. ', text)
        
        # Add space after commas if missing
        text = re.sub(r',(?=[^\s])', ', ', text)
        
        # Replace common OCR mistakes
        text = text.replace('|', 'I')  # Common OCR mistake
        text = text.replace('•', '-')  # Standardize bullet points
        text = text.replace('―', '-')  # Standardize dashes
        text = text.replace('–', '-')  # Standardize dashes
        
        # Fix spacing around bullet points
        text = re.sub(r'(?<=\w)-(?=\w)', ' - ', text)  # Add spaces around dashes between words
        
        # Remove any non-printable characters
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
        
        # Ensure sections are properly separated
        text = re.sub(r'([a-z])([A-Z][a-z])', r'\1\n\2', text)  # Add newline between lowercase and uppercase
        
        # Clean up newlines
        text = re.sub(r'\n{3,}', '\n\n', text)  # Replace multiple newlines with double newline
        
        # Normalize newlines
        text = text.replace('\r\n', '\n')
        
        # Add newlines before potential section headers
        section_headers = [
            r'EDUCATION', r'EXPERIENCE', r'SKILLS', r'CERTIFICATIONS', 
            r'PROJECTS', r'SUMMARY', r'OBJECTIVE', r'PROFILE',
            r'Education', r'Experience', r'Skills', r'Certifications',
            r'Projects', r'Summary', r'Objective', r'Profile'
        ]
        for header in section_headers:
            text = re.sub(r'([^\n])(' + header + r'[:\s])', r'\1\n\n\2', text)
        
        # Ensure bullet points are on new lines
        text = re.sub(r'([^\n])(\s*[-•*]\s+)', r'\1\n\2', text)
        
        # Fix common date formats
        text = re.sub(r'(\d{4})-(\d{1,2})-(\d{1,2})', r'\2/\3/\1', text)  # YYYY-MM-DD to MM/DD/YYYY
        
        # Add space after bullet points if missing
        text = re.sub(r'([-•*])(?=\w)', r'\1 ', text)
        
        # Ensure proper spacing for list items
        text = re.sub(r'\n([-•*])', r'\n\1', text)
        
        return text.strip()
