import io
import os
import re
import logging
from typing import Dict, List, Tuple, Optional, Any
import pdfplumber
import PyPDF2
from docx import Document
import fitz  # PyMuPDF for better PDF structure analysis
import numpy as np
from PIL import Image
import pytesseract

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)

class EnhancedDocumentProcessor:
    """Enhanced document processor that preserves document structure and formatting."""
    
    def __init__(self):
        """Initialize the document processor."""
        # Initialize OCR if available
        try:
            # Test if tesseract is available
            pytesseract.get_tesseract_version()
            self.ocr_available = True
            logger.info("Tesseract OCR is available")
        except Exception:
            self.ocr_available = False
            logger.info("Tesseract OCR is not available, image-to-text conversion will be limited")
    
    def process_file(self, file_input, file_name=None):
        """Process a document file and extract text with structure."""
        logger.info(f"Processing file input (type: {type(file_input)}), file_name: {file_name}")
        
        try:
            # Try to determine what type of input we have
            file_path = None
            
            if isinstance(file_input, str):
                if os.path.isfile(file_input):
                    file_path = file_input
                    with open(file_path, 'rb') as f:
                        file_bytes = f.read()
                    logger.info(f"Input is a file path, read {len(file_bytes)} bytes")
                else:
                    logger.warning(f"File path does not exist: {file_input}")
                    # Treat as text
                    logger.info("Treating input as plain text")
                    return self.process_plain_text(file_input)
            elif isinstance(file_input, bytes):
                file_bytes = file_input
                logger.info(f"Input is raw bytes, size: {len(file_bytes)} bytes")
            elif hasattr(file_input, 'read'):
                # File-like object
                file_bytes = file_input.read()
                if hasattr(file_input, 'name'):
                    file_path = file_input.name
                logger.info(f"Input is a file-like object, read {len(file_bytes)} bytes")
            else:
                logger.error(f"Unsupported input type: {type(file_input)}")
                return {
                    'text': '',
                    'formatted_text': '',
                    'pages': [],
                    'structure': {'sections': []},
                    'error': f"Unsupported input type: {type(file_input)}",
                    'success': False,
                    'message': "Please provide a valid file path, bytes, or file-like object."
                }
            
            # Use file_name if provided, otherwise use file_path
            file_path_for_detection = file_name if file_name else file_path
            
            # Try to detect file type from bytes
            file_type = self._detect_file_type(file_bytes, file_path_for_detection)
            logger.info(f"Detected file type: {file_type}")
            
            # Process based on file type
            if file_type == 'pdf':
                logger.info("Processing PDF file")
                result = self.process_pdf(file_bytes)
            elif file_type == 'docx':
                logger.info("Processing DOCX file")
                result = self.process_docx(file_bytes)
            elif file_type == 'image':
                logger.info("Processing image file")
                result = self.process_image(file_bytes)
            elif file_type == 'text':
                # Try to decode as text
                try:
                    text = file_bytes.decode('utf-8')
                    logger.info(f"Processing as plain text, {len(text)} characters")
                    result = self.process_plain_text(text)
                except UnicodeDecodeError:
                    logger.error("Failed to decode as UTF-8 text")
                    # Try alternative encodings
                    try:
                        text = file_bytes.decode('latin-1')
                        logger.info(f"Processing as Latin-1 text, {len(text)} characters")
                        result = self.process_plain_text(text)
                    except Exception as e:
                        logger.error(f"Failed to decode with alternative encoding: {str(e)}")
                        # Fallback to a minimal response
                        return {
                            'text': "The file contains binary data and could not be processed as a text document.",
                            'formatted_text': "The file contains binary data and could not be processed as a text document.",
                            'pages': [],
                            'structure': {'sections': []},
                            'error': "Unsupported binary format",
                            'success': False,
                            'message': "The file format is not supported. Please upload a PDF, DOCX, or TXT file."
                        }
            else:
                logger.error(f"Unsupported file type: {file_type}")
                # Try one last attempt with a simple text extraction
                try:
                    # Try to extract text using PyPDF2 in case it's a PDF with unusual signature
                    from io import BytesIO
                    from PyPDF2 import PdfReader
                    reader = PdfReader(BytesIO(file_bytes))
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n\n"
                    
                    if text.strip():
                        logger.info(f"Successfully extracted text with PyPDF2 fallback, {len(text)} characters")
                        return self.process_plain_text(text)
                except Exception as e:
                    logger.error(f"Failed PyPDF2 fallback: {str(e)}")
                
                return {
                    'text': "The file format is not supported.",
                    'formatted_text': "The file format is not supported.",
                    'pages': [],
                    'structure': {'sections': []},
                    'error': f"Unsupported file type: {file_type}",
                    'success': False,
                    'message': "The file format is not supported. Please upload a PDF, DOCX, or text file."
                }
            
            # Clean up the text and identify sections if not already done
            if 'text' in result and result['text']:
                logger.info(f"Successfully extracted text, {len(result['text'])} characters")
                result['formatted_text'] = self._clean_text(result['text'])
                
                # If no sections were identified, try to identify them based on text
                if not result.get('structure', {}).get('sections'):
                    # Add a basic structure if missing
                    if 'structure' not in result:
                        result['structure'] = {}
                    
                    # Try to identify sections by regular text patterns
                    self._identify_basic_sections(result)
                    
                    # If still no sections found and content is sufficient, create a default structure
                    if not result['structure'].get('sections') and len(result['text']) > 50:
                        # Create a minimal section structure using text-based heuristics
                        self._create_default_sections(result)
            else:
                logger.warning("No text extracted from document")
                result['error'] = "No text could be extracted from this document."
                result['success'] = False
            
            # Set success status
            result['success'] = result.get('success', True) and not result.get('error')
            return result
            
        except Exception as e:
            logger.error(f"Error in process_file: {str(e)}", exc_info=True)
            # Return a minimally valid result with error message
            return {
                'text': '',
                'formatted_text': '',
                'pages': [],
                'structure': {'sections': []},
                'error': str(e),
                'success': False,
                'message': f"Failed to process file: {str(e)}"
            }
            
    def _detect_file_type(self, file_bytes, file_path=None):
        """Detect the type of file based on content and/or path."""
        # Check file signature/magic bytes first
        if file_bytes.startswith(b'%PDF'):
            return 'pdf'
        elif file_bytes.startswith(b'PK\x03\x04'):
            return 'docx'  # Could be any zip-based format, we'll try as DOCX
        
        # Try to get file extension from path
        if file_path:
            ext = os.path.splitext(file_path)[1].lower()
            if ext in ['.pdf']:
                return 'pdf'
            elif ext in ['.docx', '.doc']:
                return 'docx'
            elif ext in ['.txt', '.text', '.md', '.rtf']:
                return 'text'
            elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
                return 'image'
        
        # Check for image file signatures
        if file_bytes.startswith(b'\x89PNG\r\n\x1a\n') or \
           file_bytes.startswith(b'\xff\xd8\xff') or \
           file_bytes.startswith(b'GIF8') or \
           file_bytes.startswith(b'II*\x00') or \
           file_bytes.startswith(b'MM\x00*'):
            return 'image'
            
        # If we can't determine, try as text
        try:
            file_bytes.decode('utf-8')
            return 'text'
        except UnicodeDecodeError:
            # If can't decode as text, return unknown
            return 'unknown'

    def _create_default_sections(self, result):
        """Create default sections based on text patterns when structure detection fails."""
        text = result.get('text', '')
        if not text:
            return
            
        sections = []
        lines = text.split('\n')
        
        # Define section keywords to look for
        section_keywords = {
            'summary': ['summary', 'profile', 'objective', 'about me', 'professional summary'],
            'experience': ['experience', 'work history', 'employment', 'professional experience', 'work experience'],
            'education': ['education', 'academic', 'qualifications', 'degrees', 'schooling'],
            'skills': ['skills', 'technical skills', 'competencies', 'expertise', 'proficiencies'],
            'certifications': ['certifications', 'certificates', 'licenses'],
            'projects': ['projects', 'project work', 'portfolio'],
            'contact': ['contact', 'contact information', 'personal information', 'phone', 'email']
        }
        
        # 1. First pass: Look for potential section headers
        potential_sections = []
        for i, line in enumerate(lines):
            line_text = line.strip()
            if not line_text:
                continue
                
            # Check if line looks like a section header (short, possibly uppercase)
            if len(line_text) < 30 and (line_text.isupper() or line_text.istitle() or line_text.endswith(':')):
                section_type = None
                
                # Check against known section types
                for section, keywords in section_keywords.items():
                    if any(keyword in line_text.lower() for keyword in keywords):
                        section_type = section
                        break
                        
                if section_type or len(line_text) < 20:  # Potential header even if section type not identified
                    potential_sections.append({
                        'index': i,
                        'text': line_text,
                        'type': section_type
                    })
        
        # 2. If no sections found this way, try more aggressive pattern matching
        if not potential_sections:
            # Look for common resume sections using regex
            for i, line in enumerate(lines):
                line_text = line.strip()
                if not line_text:
                    continue
                    
                # Check for header-like patterns with section keywords
                for section, keywords in section_keywords.items():
                    for keyword in keywords:
                        if re.search(r'\b' + re.escape(keyword) + r'\b', line_text.lower()):
                            potential_sections.append({
                                'index': i,
                                'text': line_text,
                                'type': section
                            })
                            break
                    if i in [ps['index'] for ps in potential_sections]:
                        break
        
        # 3. Create sections from identified headers
        if potential_sections:
            for i in range(len(potential_sections)):
                header = potential_sections[i]
                start_idx = header['index']
                
                # End is either the start of next section or end of document
                end_idx = potential_sections[i+1]['index'] if i < len(potential_sections) - 1 else len(lines)
                
                # Collect section content
                content_lines = lines[start_idx+1:end_idx]
                content_text = '\n'.join(line for line in content_lines if line.strip())
                
                # Create section
                section = {
                    'type': header['type'] or 'other',
                    'title': header['text'],
                    'content': content_text.strip()
                }
                
                # Only add if section has content
                if section['content']:
                    sections.append(section)
        else:
            # 4. No clear sections found - create generic sections based on text analysis
            
            # Try to detect contact info at the top
            contact_text = []
            for i in range(min(10, len(lines))):
                line = lines[i].strip()
                if re.search(r'(email|phone|address|linkedin|\w+@\w+\.\w+|\d{3}[-\.]\d{3}[-\.]\d{4})', line, re.I):
                    contact_text.append(line)
            
            if contact_text:
                sections.append({
                    'type': 'contact',
                    'title': 'Contact Information',
                    'content': '\n'.join(contact_text)
                })
            
            # Check for a summary/profile at the beginning
            summary_end = min(15, len(lines))
            has_contact = bool(contact_text)
            summary_start = len(contact_text) if has_contact else 0
            
            summary_text = '\n'.join(line for line in lines[summary_start:summary_end] if line.strip())
            if summary_text and len(summary_text) > 50:
                sections.append({
                    'type': 'summary',
                    'title': 'Professional Summary',
                    'content': summary_text
                })
            
            # Check for experience section (look for company names, job titles, dates)
            experience_indicators = r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|january|february|march|april|may|june|july|august|september|october|november|december|\d{4}[-–—]\d{4}|\d{4}[-–—]present|present)'
            experience_lines = []
            
            for i in range(summary_end, len(lines)):
                line = lines[i].strip()
                if line and re.search(experience_indicators, line, re.I):
                    # Add this line and next few lines as potential experience content
                    experience_section_end = min(i + 15, len(lines))
                    experience_lines = lines[i:experience_section_end]
                    break
                    
            if experience_lines:
                sections.append({
                    'type': 'experience',
                    'title': 'Professional Experience',
                    'content': '\n'.join(line for line in experience_lines if line.strip())
                })
            
            # Look for skills section (dense list of skills)
            skills_pattern = r'(proficient|skilled|familiar|experienced|knowledge of|expertise in|\bHTML\b|\bCSS\b|\bJava\b|\bPython\b|\bSQL\b|\bAgile\b)'
            skills_lines = []
            
            for i in range(summary_end, len(lines)):
                if not line.strip():
                    continue
                    
                if re.search(skills_pattern, line, re.I) or ('•' in line or '-' in line or ',' in line):
                    # Collect bullet point lists or comma-separated lists as potential skills
                    skills_start = i
                    skills_end = min(i + 10, len(lines))
                    skills_lines = lines[skills_start:skills_end]
                    break
                    
            if skills_lines:
                sections.append({
                    'type': 'skills',
                    'title': 'Skills',
                    'content': '\n'.join(line for line in skills_lines if line.strip())
                })
                
            # Check for education section (look for degree names, years, GPA)
            education_pattern = r'(bachelor|master|phd|doctorate|degree|university|college|school|gpa|graduated|education)'
            education_lines = []
            
            for i in range(summary_end, len(lines)):
                line = lines[i].strip()
                if line and re.search(education_pattern, line, re.I):
                    education_start = i
                    education_end = min(i + 10, len(lines))
                    education_lines = lines[education_start:education_end]
                    break
                    
            if education_lines:
                sections.append({
                    'type': 'education',
                    'title': 'Education',
                    'content': '\n'.join(line for line in education_lines if line.strip())
                })
            
        # If we still have no sections but have text, create a generic section
        if not sections and text.strip():
            sections.append({
                'type': 'general',
                'title': 'Resume Content',
                'content': text.strip()
            })
            
        # Update the result with our sections
        result['structure']['sections'] = sections
        logger.info(f"Created {len(sections)} default sections based on text analysis")
    
    def process_image(self, image_bytes: bytes) -> Dict:
        """Process an image file and extract text using OCR."""
        logger.info("Processing image file")
        
        result = {
            'pages': [{'blocks': [], 'text': '', 'page_num': 1}],
            'structure': {
                'sections': [],
                'blocks': []
            },
            'text': '',
            'formatted_text': '',
            'metadata': {},
            'from_image': True
        }
        
        try:
            # Check if OCR is available
            if not self.ocr_available:
                logger.warning("OCR is required to process images, but Tesseract is not available")
                result['text'] = "OCR is required to process images. Please install Tesseract or convert your resume to PDF, DOCX, or TXT format."
                result['formatted_text'] = result['text']
                result['pages'][0]['text'] = result['text']
                result['error'] = "OCR not available"
                return result
            
            # Process the image with OCR
            img = Image.open(io.BytesIO(image_bytes))
            
            # Convert to RGB if needed (for RGBA PNG files)
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            
            # Perform OCR
            extracted_text = pytesseract.image_to_string(img, lang='eng')
            
            if not extracted_text:
                logger.warning("OCR returned empty text")
                result['text'] = "Could not extract text from image. Please try a clearer image or convert to PDF/DOCX."
                result['formatted_text'] = result['text']
                result['pages'][0]['text'] = result['text']
                result['error'] = "OCR returned empty text"
                return result
            
            # Store extracted text
            result['text'] = extracted_text
            result['formatted_text'] = extracted_text
            result['pages'][0]['text'] = extracted_text
            
            # Add block representation
            result['pages'][0]['blocks'] = [{
                'text': extracted_text,
                'bbox': [0, 0, img.width, img.height],
                'type': 'text'
            }]
            
            # Try to identify sections in the text
            lines = extracted_text.split('\n')
            current_section = ''
            current_content = []
            sections = []
            
            for i, line in enumerate(lines):
                # Check if this looks like a section header
                if line.strip().upper() == line.strip() and len(line.strip()) > 0 and len(line.strip()) < 30:
                    # If we have content from previous section, save it
                    if current_section and current_content:
                        section_text = '\n'.join(current_content)
                        sections.append({
                            'title': current_section,
                            'content': section_text,
                            'type': self._guess_section_type(current_section, section_text)
                        })
                    
                    current_section = line.strip()
                    current_content = []
                else:
                    current_content.append(line)
            
            # Add the last section
            if current_section and current_content:
                section_text = '\n'.join(current_content)
                sections.append({
                    'title': current_section,
                    'content': section_text,
                    'type': self._guess_section_type(current_section, section_text)
                })
            
            result['structure']['sections'] = sections
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
            result['text'] = f"Error processing image: {str(e)}"
            result['formatted_text'] = result['text']
            result['pages'][0]['text'] = result['text']
            result['error'] = str(e)
            return result
    
    def _guess_section_type(self, title: str, content: str) -> str:
        """Guess the type of section based on its title and content."""
        title_lower = title.lower()
        
        if any(keyword in title_lower for keyword in ['experience', 'work', 'employment', 'career']):
            return 'experience'
        elif any(keyword in title_lower for keyword in ['education', 'academic', 'degree', 'university', 'college']):
            return 'education'
        elif any(keyword in title_lower for keyword in ['skill', 'competenc', 'proficienc', 'technical']):
            return 'skills'
        elif any(keyword in title_lower for keyword in ['summary', 'profile', 'objective', 'about']):
            return 'summary'
        elif any(keyword in title_lower for keyword in ['certification', 'certificate', 'license']):
            return 'certifications'
        elif any(keyword in title_lower for keyword in ['project', 'portfolio', 'work']):
            return 'projects'
        elif any(keyword in title_lower for keyword in ['contact', 'personal', 'info']):
            return 'contact'
        else:
            # Try to infer from content
            if re.search(r'\b(university|college|school|degree|bachelor|master|phd)\b', content.lower()):
                return 'education'
            elif re.search(r'\b(job|work|position|company|employer)\b', content.lower()):
                return 'experience'
            elif re.search(r'\b(skill|proficient|knowledge|familiar|expert)\b', content.lower()):
                return 'skills'
            
            return 'other'
    
    def process_pdf(self, file_bytes: bytes) -> Dict:
        """Process a PDF file and extract text with structure."""
        try:
            from io import BytesIO
            from PyPDF2 import PdfReader
            
            result = {
                'text': '',
                'formatted_text': '',
                'pages': [],
                'structure': {'sections': []}
            }
            
            # Create PDF reader object
            pdf = PdfReader(BytesIO(file_bytes))
            
            # Extract text from all pages
            all_text = ""
            page_texts = []
            
            # Process each page in the PDF
            for i, page in enumerate(pdf.pages):
                # Extract text from the page
                try:
                    page_text = page.extract_text()
                    if page_text:
                        all_text += page_text + "\n\n"
                        page_texts.append(page_text)
                except Exception as e:
                    logger.error(f"Error extracting text from page {i}: {str(e)}")
            
            result['text'] = all_text
            result['pages'] = page_texts
            
            # Try to identify structure
            self._identify_sections(result)
            
            return result
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            # Return a minimal valid result
            return {
                'text': '',
                'formatted_text': '',
                'pages': [],
                'structure': {'sections': []},
                'error': f"Failed to process PDF: {str(e)}"
            }
    
    def process_docx(self, file_bytes: bytes) -> Dict:
        """Process DOCX file with structure preservation."""
        result = {
            'pages': [{'blocks': [], 'text': '', 'page_num': 1}],
            'structure': {
                'sections': [],
                'blocks': []
            },
            'text': '',
            'formatted_text': '',
            'metadata': {}
        }
        
        try:
            # Load the document using bytes
            doc_io = io.BytesIO(file_bytes)
            doc = Document(doc_io)
            
            # Extract metadata if available
            core_properties = doc.core_properties
            if core_properties:
                result['metadata'] = {
                    'title': core_properties.title or '',
                    'author': core_properties.author or '',
                    'subject': core_properties.subject or '',
                    'created': str(core_properties.created) if core_properties.created else '',
                    'modified': str(core_properties.modified) if core_properties.modified else ''
                }
            
            # Process paragraphs and tables
            all_text = []
            block_num = 0
            
            for element in doc.element.body:
                # Check if it's a paragraph
                if element.tag.endswith('p'):
                    paragraph = doc.paragraphs[block_num] if block_num < len(doc.paragraphs) else None
                    if paragraph and paragraph.text.strip():
                        # Determine formatting
                        is_header = False
                        font_size = 11  # Default
                        
                        # Check if it's a heading
                        if paragraph.style and paragraph.style.name.startswith('Heading'):
                            is_header = True
                            font_size = 14  # Approximate for headings
                        
                        # Check if it has bold text
                        has_bold = any(run.bold for run in paragraph.runs if hasattr(run, 'bold'))
                        if has_bold:
                            is_header = True
                        
                        # Add to blocks
                        block_info = {
                            'text': paragraph.text,
                            'type': 'text',
                            'font_size': font_size,
                            'is_header': is_header,
                            'page': 1  # DOCX doesn't have native page info
                        }
                        
                        result['pages'][0]['blocks'].append(block_info)
                        result['structure']['blocks'].append(block_info)
                        result['pages'][0]['text'] += paragraph.text + "\n"
                        all_text.append(paragraph.text)
                
                # Check if it's a table
                elif element.tag.endswith('tbl'):
                    for table in doc.tables:
                        table_text = []
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                table_text.append(' | '.join(row_text))
                        
                        if table_text:
                            block_text = '\n'.join(table_text)
                            block_info = {
                                'text': block_text,
                                'type': 'table',
                                'font_size': 11,  # Default
                                'is_header': False,
                                'page': 1
                            }
                            
                            result['pages'][0]['blocks'].append(block_info)
                            result['structure']['blocks'].append(block_info)
                            result['pages'][0]['text'] += block_text + "\n\n"
                            all_text.append(block_text)
                
                block_num += 1
            
            # Set the full text
            result['text'] = "\n".join(all_text)
            
            # Create formatted text with section markers
            formatted_lines = []
            prev_is_header = False
            
            for block in result['structure']['blocks']:
                if block['type'] == 'text':
                    text = block['text'].strip()
                    if not text:
                        continue
                    
                    # Add extra newlines before headers
                    if block['is_header'] and not prev_is_header:
                        formatted_lines.append("")  # Extra newline
                    
                    formatted_lines.append(text)
                    prev_is_header = block['is_header']
                
                elif block['type'] == 'table':
                    formatted_lines.append("")  # Extra newline before table
                    formatted_lines.append(block['text'])
                    formatted_lines.append("")  # Extra newline after table
                    prev_is_header = False
            
            result['formatted_text'] = "\n".join(formatted_lines)
            
            # Identify document sections
            self._identify_sections(result)
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
        
        return result
    
    def process_plain_text(self, text: str) -> Dict:
        """Process plain text with basic structure detection."""
        result = {
            'pages': [{'blocks': [], 'text': text, 'page_num': 1}],
            'structure': {
                'sections': [],
                'blocks': []
            },
            'text': text,
            'formatted_text': '',
            'metadata': {}
        }
        
        # Split into lines
        lines = text.split('\n')
        
        # Heuristic to identify headers: all caps, short lines, etc.
        blocks = []
        current_block = {'text': '', 'lines': [], 'is_header': False}
        
        for line in lines:
            line = line.rstrip()
            
            # Empty line marks a block boundary
            if not line.strip():
                if current_block['lines']:
                    block_text = '\n'.join(current_block['lines'])
                    blocks.append({
                        'text': block_text,
                        'type': 'text',
                        'is_header': current_block['is_header'],
                        'font_size': 14 if current_block['is_header'] else 11,  # Approximation
                        'page': 1
                    })
                    current_block = {'text': '', 'lines': [], 'is_header': False}
                continue
            
            # Check if line looks like a header
            is_line_header = False
            if line.isupper() and len(line) < 50:
                is_line_header = True
            elif re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+){0,3}$', line) and len(line) < 30:
                is_line_header = True
            elif re.match(r'^[A-Z][A-Za-z]*\s*:$', line):
                is_line_header = True
            
            # If current line is a header and we have content, finish current block
            if is_line_header and current_block['lines']:
                block_text = '\n'.join(current_block['lines'])
                blocks.append({
                    'text': block_text,
                    'type': 'text',
                    'is_header': current_block['is_header'],
                    'font_size': 14 if current_block['is_header'] else 11,
                    'page': 1
                })
                current_block = {'text': '', 'lines': [], 'is_header': is_line_header}
            
            # Update current block
            if is_line_header:
                current_block['is_header'] = True
            
            current_block['lines'].append(line)
        
        # Add the last block if it exists
        if current_block['lines']:
            block_text = '\n'.join(current_block['lines'])
            blocks.append({
                'text': block_text,
                'type': 'text',
                'is_header': current_block['is_header'],
                'font_size': 14 if current_block['is_header'] else 11,
                'page': 1
            })
        
        # Update result with blocks
        result['structure']['blocks'] = blocks
        result['pages'][0]['blocks'] = blocks
        
        # Create formatted text
        formatted_lines = []
        for block in blocks:
            if block['is_header']:
                formatted_lines.append("")  # Add newline before header
            
            formatted_lines.append(block['text'])
        
        result['formatted_text'] = "\n".join(formatted_lines)
        
        # Identify sections
        self._identify_basic_sections(result)
        
        return result
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text while preserving structure."""
        if not text:
            return ""
        
        # Replace excessive newlines with single newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Ensure section headers are properly separated
        section_headers = [
            r'EDUCATION', r'EXPERIENCE', r'SKILLS', r'CERTIFICATIONS', 
            r'PROJECTS', r'SUMMARY', r'OBJECTIVE', r'PROFILE',
            r'Education', r'Experience', r'Skills', r'Certifications',
            r'Projects', r'Summary', r'Objective', r'Profile'
        ]
        
        for header in section_headers:
            text = re.sub(r'([^\n])(' + header + r'[:\s])', r'\1\n\n\2', text)
        
        # Ensure bullet points are on new lines
        text = re.sub(r'([^\n])(\s*[-•*]\s+)', r'\1\n\2', text)
        
        # Fix common date formats for better parsing
        text = re.sub(r'(\d{4})-(\d{1,2})-(\d{1,2})', r'\2/\3/\1', text)  # YYYY-MM-DD to MM/DD/YYYY
        
        # Add space after bullet points if missing
        text = re.sub(r'([-•*])(?=\w)', r'\1 ', text)
        
        # Ensure proper spacing for list items
        text = re.sub(r'\n([-•*])', r'\n\1', text)
        
        return text.strip()
    
    def _identify_sections(self, result):
        """Identify and extract sections from document structure."""
        try:
            # Get formatted text and blocks
            formatted_text = result.get('formatted_text', '')
            blocks = result.get('structure', {}).get('blocks', [])
            
            if not formatted_text or not blocks:
                logger.warning("No formatted text or blocks found for section identification")
                return
                
            # Common section header patterns
            section_patterns = {
                'summary': r'(?i)^(?:summary|profile|professional\s+summary|career\s+summary|objective|about\s+me)$',
                'experience': r'(?i)^(?:experience|work\s+experience|professional\s+experience|employment\s+history|work\s+history)$',
                'education': r'(?i)^(?:education|academic\s+background|qualifications|academic\s+qualifications)$',
                'skills': r'(?i)^(?:skills|technical\s+skills|core\s+competencies|expertise|proficiencies)$',
                'certifications': r'(?i)^(?:certifications|certificates|professional\s+certifications|accreditations)$',
                'projects': r'(?i)^(?:projects|personal\s+projects|portfolio|project\s+experience)$'
            }
            
            # Try to identify sections based on formatting and text patterns
            potential_headers = []
            
            # First pass: identify potential section headers
            for i, block in enumerate(blocks):
                if block.get('type') != 'text':
                    continue
                    
                text = block.get('text', '').strip()
                if not text:
                    continue
                    
                # Check if this could be a section header based on formatting
                is_header = block.get('is_header', False) or (block.get('font_size', 0) > 11)
                
                # Check if text matches any section pattern
                section_type = None
                for section, pattern in section_patterns.items():
                    if re.match(pattern, text, re.IGNORECASE):
                        section_type = section
                        break
                        
                if is_header or section_type:
                    potential_headers.append({
                        'index': i,
                        'text': text,
                        'section_type': section_type
                    })
            
            # If no headers were found using formatting, try based on text patterns alone
            if not potential_headers:
                # Split formatted text into lines
                lines = formatted_text.split('\n')
                line_indices = {}
                
                # Map line text to their positions in the blocks
                for i, line in enumerate(lines):
                    line_text = line.strip()
                    if not line_text:
                        continue
                        
                    for j, block in enumerate(blocks):
                        if block.get('type') != 'text':
                            continue
                            
                        block_text = block.get('text', '').strip()
                        if line_text in block_text:
                            line_indices[i] = j
                            break
                
                # Check each line for section headers
                for i, line in enumerate(lines):
                    line_text = line.strip()
                    if not line_text:
                        continue
                        
                    # Check if this line matches any section pattern
                    section_type = None
                    for section, pattern in section_patterns.items():
                        if re.match(pattern, line_text, re.IGNORECASE):
                            section_type = section
                            break
                            
                    if section_type and i in line_indices:
                        potential_headers.append({
                            'index': line_indices[i],
                            'text': line_text,
                            'section_type': section_type
                        })
            
            # If we still have no headers, use heuristic approach
            if not potential_headers:
                potential_headers = self._identify_sections_heuristic(blocks)
            
            # Second pass: create sections from headers
            sections = []
            
            if potential_headers:
                for i in range(len(potential_headers)):
                    header = potential_headers[i]
                    start_idx = header['index']
                    
                    # End is either the start of next section or end of document
                    end_idx = potential_headers[i+1]['index'] if i < len(potential_headers) - 1 else len(blocks)
                    
                    # Skip if start and end are the same
                    if start_idx >= end_idx:
                        continue
                    
                    # Determine section type
                    section_type = header['section_type']
                    
                    # If section type isn't known, try to guess based on content
                    if not section_type:
                        section_type = self._guess_section_type(header['text'])
                    
                    # Collect content of section
                    content_blocks = blocks[start_idx+1:end_idx]
                    content_text = ""
                    
                    for block in content_blocks:
                        if block.get('type') == 'text':
                            content_text += block.get('text', '') + "\n"
                    
                    # Trim trailing whitespace
                    content_text = content_text.strip()
                    
                    # Create section
                    section = {
                        'type': section_type or 'other',
                        'title': header['text'],
                        'content': content_text,
                        'start_idx': start_idx,
                        'end_idx': end_idx
                    }
                    
                    sections.append(section)
            
            # Store sections in result
            result['structure']['sections'] = sections
            logger.info(f"Identified {len(sections)} sections in document")
            
        except Exception as e:
            logger.error(f"Error identifying sections: {str(e)}")
            # Set empty sections list as fallback
            result['structure']['sections'] = []
    
    def _identify_sections_heuristic(self, blocks):
        """Use heuristics to identify section headers when standard methods fail."""
        potential_headers = []
        
        for i, block in enumerate(blocks):
            if block.get('type') != 'text':
                continue
                
            text = block.get('text', '').strip()
            if not text:
                continue
                
            # Check for short lines that might be headers
            is_short = len(text) < 50
            
            # Check if text is all caps or title case
            is_emphasized = text.isupper() or text.istitle()
            
            # Check if isolated by spacing (preceded or followed by empty blocks)
            is_isolated = False
            if i > 0 and i < len(blocks) - 1:
                prev_block = blocks[i-1]
                next_block = blocks[i+1]
                
                prev_text = prev_block.get('text', '').strip() if prev_block.get('type') == 'text' else ''
                next_text = next_block.get('text', '').strip() if next_block.get('type') == 'text' else ''
                
                is_isolated = not prev_text or not next_text
            
            # If likely a header, add to list
            if is_short and (is_emphasized or is_isolated):
                potential_headers.append({
                    'index': i,
                    'text': text,
                    'section_type': None  # Will be determined later
                })
        
        return potential_headers
    
    def _guess_section_type(self, text):
        """Guess section type based on header text."""
        text_lower = text.lower()
        
        if any(term in text_lower for term in ['summary', 'profile', 'objective', 'about']):
            return 'summary'
        elif any(term in text_lower for term in ['experience', 'work', 'employment', 'job', 'career']):
            return 'experience'
        elif any(term in text_lower for term in ['education', 'academic', 'degree', 'university', 'school']):
            return 'education'
        elif any(term in text_lower for term in ['skill', 'competency', 'proficiency', 'qualification']):
            return 'skills'
        elif any(term in text_lower for term in ['certification', 'certificate', 'license']):
            return 'certifications'
        elif any(term in text_lower for term in ['project', 'portfolio', 'work sample']):
            return 'projects'
        else:
            return 'other'
    
    def _identify_basic_sections(self, result: Dict) -> None:
        """Basic section identification for fallback methods."""
        if not result or 'text' not in result:
            logging.warning("No text available for section identification")
            return
            
        text = result['text']
        lines = text.split('\n')
        
        # Enhanced section patterns with more variations and better boundaries
        section_patterns = {
            'summary': r'(?i)^(?:summary|profile|objective|about|professional\s+summary|career\s+objective)(?:\s*:)?$',
            'education': r'(?i)^(?:education|academic|qualification|degree|educational\s+background)(?:\s*:)?$',
            'experience': r'(?i)^(?:experience|employment|work\s+history|professional\s+experience|work\s+experience)(?:\s*:)?$',
            'skills': r'(?i)^(?:skills|technical\s+skills|expertise|competencies|core\s+competencies)(?:\s*:)?$',
            'projects': r'(?i)^(?:projects|personal\s+projects|academic\s+projects|project\s+experience)(?:\s*:)?$',
            'certifications': r'(?i)^(?:certifications|certificates|accreditations|professional\s+certifications)(?:\s*:)?$',
            'awards': r'(?i)^(?:awards|honors|achievements|recognition)(?:\s*:)?$',
            'languages': r'(?i)^(?:languages|language\s+proficiency|language\s+skills)(?:\s*:)?$',
            'interests': r'(?i)^(?:interests|hobbies|activities|personal\s+interests)(?:\s*:)?$',
            'references': r'(?i)^(?:references|recommendations|professional\s+references)(?:\s*:)?$',
            'contact': r'(?i)^(?:contact|contact\s+information|personal\s+information)(?:\s*:)?$'
        }
        
        # First pass: Find all potential section headers
        section_starts = []
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Check if the line is a section header
            for section, pattern in section_patterns.items():
                if re.match(pattern, line, re.IGNORECASE) or (
                    len(line) < 25 and  # Short line, likely a header
                    re.search(pattern, line, re.IGNORECASE) and 
                    all(c.isupper() for c in line if c.isalpha()) or  # All uppercase
                    line.endswith(':')  # Ends with colon
                ):
                    section_starts.append((i, section, line))
                    logging.info(f"Found section '{section}' at line {i}: '{line}'")
                    break

        # Second pass: Check for formatting patterns (all caps, bold, etc.)
        if not section_starts:
            logging.warning("No section headers found with pattern matching, trying formatting heuristics")
            for i, line in enumerate(lines):
                line = line.strip()
                if not line or len(line) > 30:  # Skip empty lines or long lines
                    continue
                    
                # Check for all uppercase or title case for potential headers
                if all(c.isupper() for c in line if c.isalpha()) or line.istitle():
                    for section, pattern in section_patterns.items():
                        section_keyword = re.sub(r'(?:\^|\(\?i\)|\(|\)|\\s\*|\?|:|\$)', '', pattern).split('|')[0]
                        if section_keyword.lower() in line.lower():
                            section_starts.append((i, section, line))
                            logging.info(f"Found section '{section}' using formatting at line {i}: '{line}'")
                            break
        
        # Create sections with better boundary detection
        sections = []
        if section_starts:
            # Sort by line index
            section_starts.sort(key=lambda x: x[0])
            
            # Determine section boundaries
            for i, (line_idx, section_type, header) in enumerate(section_starts):
                start_idx = line_idx
                
                # Determine end index (next section or end of document)
                if i < len(section_starts) - 1:
                    end_idx = section_starts[i + 1][0]
                else:
                    end_idx = len(lines)
                
                # Extract section content, skipping the header
                content_start = start_idx + 1
                while content_start < end_idx and not lines[content_start].strip():
                    content_start += 1  # Skip empty lines after header
                
                content_end = end_idx
                while content_end > content_start and not lines[content_end - 1].strip():
                    content_end -= 1  # Skip empty lines at the end
                
                # Get section content with proper handling of empty sections
                if content_start < content_end:
                    content = '\n'.join(lines[content_start:content_end])
                else:
                    content = ""  # Empty section
                
                # Check if the content is meaningful
                # Sometimes we get false positives where a word matches a section pattern
                if content and len(content) < 10 and content_end - content_start <= 1:
                    continue  # Skip likely false positives
                
                section = {
                    'type': section_type,
                    'header': header,
                    'content': content,
                    'start_index': start_idx,
                    'end_index': end_idx,
                    'char_start': sum(len(line) + 1 for line in lines[:start_idx]),
                    'char_end': sum(len(line) + 1 for line in lines[:end_idx]) 
                }
                sections.append(section)
                logging.info(f"Created section '{section_type}' from line {start_idx} to {end_idx} with {len(content)} chars")
        
        # If no sections found, create artificial ones based on content analysis
        if not sections:
            logging.warning("No sections identified, creating artificial sections")
            
            # Look for contact info at the top
            contact_end = min(10, len(lines) // 4)  # First 10 lines or first quarter
            
            # Email pattern
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            phone_pattern = r'\b(?:\+?\d{1,2}\s?)?(?:\(?\d{3}\)?[\s.-]?)?\d{3}[\s.-]?\d{4}\b'
            
            # Check for contact info in the first few lines
            found_contact = False
            for i in range(min(contact_end, len(lines))):
                if re.search(email_pattern, lines[i]) or re.search(phone_pattern, lines[i]):
                    found_contact = True
                    sections.append({
                        'type': 'contact',
                        'header': 'Contact Information',
                        'content': '\n'.join(lines[:contact_end]),
                        'start_index': 0,
                        'end_index': contact_end,
                        'char_start': 0,
                        'char_end': sum(len(line) + 1 for line in lines[:contact_end])
                    })
                    break
            
            # Create sections based on content analysis
            remaining_text = '\n'.join(lines[contact_end if found_contact else 0:])
            
            # Try to identify a summary/objective section
            if len(remaining_text) > 100:
                summary_end = len(remaining_text) // 5  # First 20% might be summary
                sections.append({
                    'type': 'summary',
                    'header': 'Summary',
                    'content': remaining_text[:summary_end],
                    'start_index': contact_end if found_contact else 0,
                    'end_index': contact_end + summary_end if found_contact else summary_end,
                    'char_start': sum(len(line) + 1 for line in lines[:contact_end if found_contact else 0]),
                    'char_end': sum(len(line) + 1 for line in lines[:contact_end if found_contact else 0]) + summary_end
                })
                
                # Add a generic experience section for the rest
                sections.append({
                    'type': 'experience',
                    'header': 'Experience',
                    'content': remaining_text[summary_end:],
                    'start_index': (contact_end if found_contact else 0) + summary_end,
                    'end_index': len(lines),
                    'char_start': sum(len(line) + 1 for line in lines[:(contact_end if found_contact else 0)]) + summary_end,
                    'char_end': sum(len(line) + 1 for line in lines)
                })
            else:
                # Just one generic section
                sections.append({
                    'type': 'content',
                    'header': 'Content',
                    'content': remaining_text,
                    'start_index': contact_end if found_contact else 0,
                    'end_index': len(lines),
                    'char_start': sum(len(line) + 1 for line in lines[:contact_end if found_contact else 0]),
                    'char_end': sum(len(line) + 1 for line in lines)
                })
        
        result['structure']['sections'] = sections
    
    def _perform_ocr(self, file_bytes: bytes) -> str:
        """Perform OCR on the document if necessary."""
        try:
            # Convert PDF to images
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_parts = []
            
            for page_num, page in enumerate(doc):
                # Convert page to image
                pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Perform OCR
                text = pytesseract.image_to_string(img, lang='eng')
                if text:
                    text_parts.append(text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            logger.error(f"OCR failed: {str(e)}")
            return "" 